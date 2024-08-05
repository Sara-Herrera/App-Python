import logging
import os
import json
import re
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox, font, filedialog
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import asksaveasfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import pandas as pd
from PIL import Image, ImageTk
from tkcalendar import DateEntry
import babel.numbers
import locale
from pypdf import PdfReader
from tkcalendar import Calendar

class InterfazApp:
    CONFIG_FILE_NAME = "config.json"
    PATH_LABELS = "labels.py"
    FILE_PATH = ""
    STUDY_PATH = ""
    FONT_NAME = "",
    APP_VERSION = ""
    

    def __init__(self, root):
        self.app_title = None
        self.root = root
        self.read_config()
        self.init_window()
        self.setup_fonts_and_frames()
        self.root.mainloop()

    # Método que extrae información del archivo de configuración y carga datos en variables globales
    def read_config(self):
        # Obtener la ruta completa del archivo de configuración utilizando resource_path
        config_file_path = self.resource_path(InterfazApp.CONFIG_FILE_NAME)
        self.PATH_LABELS = self.resource_path(InterfazApp.PATH_LABELS)

        if not os.path.exists(config_file_path):
            messagebox.showerror("Error", "No se localiza el fichero de configuración")
            logging.debug("Creando fichero de configuración")
        logging.debug("Leyendo fichero de configuración")
        with open(config_file_path, mode="r") as json_file:
            config = json.load(json_file)
        self.app_title = config["APP_TITLE"]
        self.window_width = int(config["WINDOW_WIDTH"])
        self.window_height = int(config["WINDOW_HEIGHT"])
        self.FONT_NAME = config["FONT_NAME"]
        self.APP_VERSION = config["APP_VERSION"]

    # Método que inicializa la interfaz con el tamaño indicado
    def init_window(self):
        logging.debug("Entering init_window")
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        OFFSET_X = int((screen_width - self.window_width) / 2)
        OFFSET_Y = int((screen_height - self.window_height) / 2)
        self.root.geometry(f"{self.window_width}x{self.window_height}+{OFFSET_X}+{OFFSET_Y}")
        self.root.resizable(False, False)
        self.root.title(self.app_title)

    # Método que establece componentes de interfaz principal
    def setup_fonts_and_frames(self):
        self.font_disclaimer = font.Font(family=self.FONT_NAME, size=9)
        self.font_buttom_global = font.Font(family=self.FONT_NAME, size=10, weight=font.BOLD)
        self.font_etiquetas = font.Font(family=self.FONT_NAME, size=11, weight=font.BOLD)
        self.font_buttom_specific = font.Font(family=self.FONT_NAME, size=10, weight=font.BOLD)
        self.font_derechos = font.Font(family=self.FONT_NAME, size=7, weight=font.BOLD)

        self.image_background_top = ImageTk.PhotoImage(Image.open(self.resource_path('images_folder\\fondo.jpg')), master=root)

        # Frame superior con botones
        self.frame_top = tk.Frame(self.root)
        self.frame_top.pack(side=tk.TOP, fill=tk.BOTH, pady=3)

        self.background_top = tk.Label(self.frame_top, image = self.image_background_top)
        self.background_top.place(relwidth=1, relheight=1)

        self.buttom_study = tk.Button(self.frame_top, text="CREAR ARCHIVO\n DATOS DE ESTUDIO", font=self.font_buttom_global, height=3, width=20, background="#ccada3", fg='#3c3c3c', command=self.view_create_study_data)
        self.buttom_study.grid(row=0, column=0, padx=200, pady=0)

        self.buttom_report = tk.Button(self.frame_top, text="CREAR INFORME", font=self.font_buttom_global, height=3, width=20, background="#ccada3", fg='#3c3c3c', command=self.view_generate_report)
        self.buttom_report.grid(row=0, column=1, padx=0, pady=0)
        
        # Frame derecho para mostrar componentes según el botón seleccionado
        self.frame_bottom = tk.Frame(self.root)
        self.frame_bottom.pack(side=tk.BOTTOM, expand=True, fill=tk.BOTH)
        
        self.background_bottom = tk.Label(self.frame_bottom, background = "#ffd9cc")
        self.background_bottom.place(relwidth=1, relheight=1, relx=0.5, rely=0.5, anchor='center')

        # Crear un Text widget con scrollbar
        self.text_frame = tk.Frame(self.frame_bottom)
        self.text_frame.pack(expand=True, fill=tk.BOTH, padx=20, pady=50)
        
        self.text_widget = tk.Text(self.text_frame, wrap='word',font=self.font_disclaimer, fg='#3c3c3c', bg=self.frame_bottom.cget('bg'))
        self.text_widget.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        
        self.scrollbar = ttk.Scrollbar(self.text_frame, orient=tk.VERTICAL, command=self.text_widget.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_widget.config(yscrollcommand=self.scrollbar.set)
        
        # Función para cargar el texto en el Text widget
        self.text_widget.insert(tk.END, self.load_labels(self.PATH_LABELS)['texto_disclaimer'])
        self.text_widget.config(state=tk.DISABLED)
        
        self.label_disclaimer = tk.Label(self.frame_bottom, text=f"{self.app_title} {self.APP_VERSION} | Todos los derechos reservados © 2024 \n Powered by Sara Herrera", font=self.font_derechos, fg='#3c3c3c', bg='#ffd9cc')
        self.label_disclaimer.pack(padx=0, pady=5)
 
    # Método que permite cambiar la parte inferior de la ventana
    def clear_frame_bottom(self):
        # Limpiar el frame inferior eliminando todos los componentes
        for widget in self.frame_bottom.winfo_children():
            widget.destroy()

    # Método para establecer la ruta interna
    def resource_path(self, relative_path):
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)
    
    # Método que carga la información de las etiquetas
    def load_labels(self, labels_path):
        labels = {}
    
        labels_file_path = self.resource_path(labels_path)
        
        with open(labels_file_path, 'r', encoding='utf-8') as f:
            exec(f.read(), labels)
    
        return labels
            
# -----------------------------------------
#     VISTA CREACIÓN DE DATOS DE ESTUDIO
# -----------------------------------------    
    
    # Método para establecer los componentes de la interfaz de datos de estudio y dar funcionalidad a los botones
    def view_create_study_data(self):
        # Limpiar el frame antes de agregar nuevos componentes
        self.clear_frame_bottom()

        s = ttk.Style()
        # Configurar los estilos
        s.configure("Custom.Treeview.Heading", font=(self.FONT_NAME, 10, 'bold'))
        s.configure("Custom.Treeview", font=(self.FONT_NAME, 10), rowheight=35)
        s.configure('style_buttom_specific.TButton', font=(self.FONT_NAME, 12), background='#ffd9cc', foreground='#3c3c3c')

        # Fondo del frame
        self.background_bottom = tk.Label(self.frame_bottom, background="#ffd9cc")
        self.background_bottom.place(relwidth=1, relheight=1)

        # Crear TreeView
        self.tree = ttk.Treeview(self.frame_bottom, selectmode='browse')
        self.tree.grid(row=0, column=0, columnspan=4, sticky='nsew', padx=210, pady=5)
        self.tree["columns"] = ("0", "1", "2", "3")
        self.tree['show'] = 'headings'

        # Definir el ancho de las columnas
        self.tree.column("0", width=80, anchor='c')
        self.tree.column("1", width=130, anchor='c')
        self.tree.column("2", width=180, anchor='c')
        self.tree.column("3", width=125, anchor='c')

        # Definir el nombre de las cabeceras
        self.tree.heading("0", text="Estudio ID")
        self.tree.heading("1", text="Objetivo")
        self.tree.heading("2", text="Enfermedad/Gen de estudio")
        self.tree.heading("3", text="Filtros previos")

        self.scrollbar = ttk.Scrollbar(self.frame_bottom, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscroll=self.scrollbar.set)
        self.scrollbar.grid(row=0, column=3, sticky='ns')

        # Crear componentes para entrada de datos
        # Estudio ID
        self.label_study = tk.Label(self.frame_bottom, text='Estudio ID', width=15, font=(self.FONT_NAME, 12), bg='#ffd9cc')
        self.label_study.grid(row=1, column=1, pady=10)
        self.text_study = tk.Entry(self.frame_bottom)
        self.text_study.grid(row=1, column=2, sticky='w')

        # Objetivo
        self.label_objective = tk.Label(self.frame_bottom, text='Objetivo', width=15, font=(self.FONT_NAME, 12), bg='#ffd9cc')
        self.label_objective.grid(row=2, column=1, pady=10)
        self.text_objective = tk.Entry(self.frame_bottom)
        self.text_objective.grid(row=2, column=2, sticky='w')

        # Enfermedad / gen
        self.label_disease = tk.Label(self.frame_bottom, text='Enfermedad/Gen', width=15, font=(self.FONT_NAME, 12), bg='#ffd9cc')
        self.label_disease.grid(row=3, column=1, pady=10)
        self.text_disease = tk.Entry(self.frame_bottom)
        self.text_disease.grid(row=3, column=2, sticky='w')

        # Filtro
        self.filter_res = tk.StringVar(value='No informacion')  # Definir el tipo de variable para los valores del botón BCL6
        self.filter = tk.Label(self.frame_bottom, text='Filtros previos', width=15, font=(self.FONT_NAME, 12), bg='#ffd9cc')
        self.filter.grid(row=4, column=0, padx = 0, pady = 5)
        self.filter_1 = tk.Radiobutton(self.frame_bottom, text='Si', variable=self.filter_res, value='Si', bg='#ffd9cc')
        self.filter_1.grid(row=4, column=1, padx = 0, pady = 5, sticky='w')
        self.filter_0 = tk.Radiobutton(self.frame_bottom, text='No', variable=self.filter_res, value='No', bg='#ffd9cc')
        self.filter_0.grid(row=4, column=2, padx = 0, pady = 5, sticky='w')
        self.filter_2 = tk.Radiobutton(self.frame_bottom, text='No informacion', variable=self.filter_res, value='No informacion', bg='#ffd9cc')
        self.filter_2.grid(row=4, column=3, padx = 0, pady = 5, sticky='w')

        # Botones
        self.button_insert = ttk.Button(self.frame_bottom, text='INSERTAR', style='style_buttom_specific.TButton', command=self.insert)
        self.button_insert.grid(row = 6, column = 1, padx = 0, pady = 5)

        self.button_reset = ttk.Button(self.frame_bottom, text='RESETEAR', style='style_buttom_specific.TButton', command=self.reset)
        self.button_reset.grid(row = 6, column = 2, padx = 0, pady = 5)

        self.button_delete = ttk.Button(self.frame_bottom, text='BORRAR', style='style_buttom_specific.TButton', command=self.delete)
        self.button_delete.grid(row = 7, column = 2, padx = 0, pady = 5)

        self.button_export = ttk.Button(self.frame_bottom, text='EXPORTAR', style='style_buttom_specific.TButton', command=self.export)
        self.button_export.grid(row = 7, column = 1, padx = 0, pady = 5)

    # Método para insertar fila en la tabla
    def insert(self):
        study = self.text_study.get()
        objective = self.text_objective.get()
        disease = self.text_disease.get()
        filter_res = self.filter_res.get()

        self.tree.insert("", 'end', values = (study, objective, disease, filter_res))

    # Método para borrar la tabla
    def reset(self):
        reseted = messagebox.askokcancel("Eliminar", "¿Seguro que quieres eliminar la tabla?")
        if reseted:
            for item in self.tree.get_children():
                self.tree.delete(item)
        else:
            return

    # Método para borrar fila en la tabla
    def delete(self):
        deleted = messagebox.askokcancel("Eliminar", "¿Seguro que quieres eliminar la fila?")
        if deleted:
            selected_item = self.tree.selection()[0]
            self.tree.delete(selected_item)
        else:
            return
        
    # Método para exportar la tabla
    def export(self):
        try:
            path = asksaveasfile(mode='wb', defaultextension=".txt")
        except:
            return

        row_list =[]
        columns = ('Estudio.ID', 'Objetivo', 'Enfermedad/Gen', 'Filtro')
        for child in self.tree.get_children():
            row_list.append(self.tree.item(child)['values'])
        tree_df = pd.DataFrame(row_list, columns = columns)
        tree_df.to_csv(path, header=True, index=False, sep='\t', mode='a')

        save = messagebox.showinfo('Guardar', 'Se ha generado el archivo correctamente. ')
      
    
# -----------------------------------------
#         VISTA CREACIÓN DE INFORME
# -----------------------------------------       
    
    # Método para establecer los componentes de la interfaz de creación de informe y dar funcionalidad a los botones
    def view_generate_report(self):
        
        s = ttk.Style()
        
        # Configurar los estilos
        s.configure('style_buttom_specific.TButton', font=(self.FONT_NAME, 12), background='#ffd9cc', foreground='#3c3c3c')
        
        # Limpiar el frame antes de agregar nuevos componentes
        self.clear_frame_bottom()
        
        # Fondo del frame
        self.background_bottom = tk.Label(self.frame_bottom, background = "#ffd9cc")
        self.background_bottom.place(relwidth=1, relheight=1)

        #Inicializar componentes de la interfaz
        # Fila 1
        label_cargar = ttk.Label(self.frame_bottom, text="Cargar archivos:", font=self.font_etiquetas, background="#ffd9cc", foreground = "#3c3c3c")
        label_cargar.grid(row=1, column=1, padx=190, pady=30)
        
        # Fila 3
        button_load_analysis_data = ttk.Button(self.frame_bottom, text="Datos de variantes", style='style_buttom_specific.TButton', command=self.load_data)
        button_load_analysis_data.grid(row=3, column=1, padx = 0, pady = 3)

        buttom_load_study_data = ttk.Button(self.frame_bottom, text="Datos de estudio ", style='style_buttom_specific.TButton', command=self.load_study_data)
        buttom_load_study_data.grid(row=3, column=2, padx = 0, pady = 3)
 
        # Fila 4
        self.label_load_analysis_data = ttk.Label(self.frame_bottom, text="", font=self.font_etiquetas, background="#ffd9cc", foreground = "#737373")
        self.label_load_analysis_data.grid(row=4, column=1, padx = 0, pady = 20)
        
        self.label_load_study_data = ttk.Label(self.frame_bottom, text="", font=self.font_etiquetas, background="#ffd9cc", foreground = "#737373")
        self.label_load_study_data.grid(row=4, column=2, padx = 0, pady = 20)        

        # Fila 6
        label_id_study = ttk.Label(self.frame_bottom, text="ID del estudio:", font=self.font_etiquetas, background="#ffd9cc", foreground = "#3c3c3c")
        label_id_study.grid(row=6, column=1, padx = 0, pady=30)
        
        text_id_study = tk.Entry(self.frame_bottom)
        text_id_study.grid(row=6, column=2, padx = 0, pady=30)
        
        # Fila 7
        label_1 = ttk.Label(self.frame_bottom, text="Filtrar por Germline Classification:", font=self.font_etiquetas, background="#ffd9cc", foreground = "#3c3c3c")
        label_1.grid(row=7, column=1, padx = 10, pady=10)

        self.combo_1 = ttk.Combobox(self.frame_bottom)
        self.combo_1.grid(row=7, column=2, padx = 0, pady=10, ipady=0, ipadx=0)
        self.combo_1.set("")
        self.combo_1.config(postcommand=lambda: self.update_cmb(self.combo_1))
        
        # Fila 8
        label_fecha = ttk.Label(self.frame_bottom, text="Fecha:", font=self.font_etiquetas, background="#ffd9cc", foreground = "#3c3c3c")
        label_fecha.grid(row=8, column=1, padx = 0, pady=30)
        
        fecha_seleccionada = DateEntry(self.frame_bottom, selectmode = "day", locale='es_ES')
        fecha_seleccionada.grid(row=8, column=2, padx = 0, pady=30) 
        
        # Fila 10
        buttom_generate = ttk.Button(self.frame_bottom, text="GENERAR", style='style_buttom_specific.TButton', command=lambda: [self.generate_report(self.combo_1.get(), fecha_seleccionada.get(), text_id_study.get()), self.clear_view(self.combo_1.get(), fecha_seleccionada, text_id_study)])
        buttom_generate.grid(row=10, column=1, columnspan=2, padx = 0, pady=30, ipady=20, ipadx=50)
        
        # Agregar valores al combobox
        self.load_cmb_values(self.combo_1)

        
    # Método para cargar los valores en el combobox
    def load_cmb_values(self, combobox):
        valores = ["All", "Conflicting classifications", "Benign", "Likely benign", "Uncertain significance", "Likely pathogenic", "Pathogenic"]

        # Limpiar y llenar el combobox con los nuevos valores
        combobox['values'] = valores
        combobox.set("")  # Limpiar la selección actual si la hay
            
        return valores
         
    # Método para actualizar valores en el combobox si se ha escrito texto para realizar una búsqueda
    def update_cmb(self, combobox):
        valor_escrito = combobox.get()
        valores_totales = self.load_cmb_values(combobox)
        valores_filtrados = [valor for valor in valores_totales if valor_escrito.lower() in valor.lower()]
        combobox['values'] = valores_filtrados
        combobox.set(valor_escrito)
        
    # Método para cargar el archivo de datos y establecer su ruta en variable global        
    def load_data(self):
        file_path = askopenfilename(title="Seleccionar archivo", filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")])
        if file_path:
            self.FILE_PATH = file_path
            self.label_load_analysis_data.configure(text="Archivo cargado correctamente. ")

    # Método para cargar el archivo de datos de estudio y establecer su ruta en variable global        
    def load_study_data(self):
        file_path = askopenfilename(title="Seleccionar archivo", filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")])
        if file_path:
            self.STUDY_PATH = file_path
            self.label_load_study_data.configure(text="Archivo cargado correctamente. ")
            
    # Método para limpiar variables y rutas de archivos de interfaz tras generar informe
    def clear_view(self, combo_1, fecha_seleccionada, text_id_study):
        self.combo_1.set('')  # Limpiar valor seleccionado en el combobox
        fecha_actual = datetime.today().strftime('%Y-%m-%d')
        fecha_seleccionada.delete(0, tk.END)  # Limpiar entrada de fecha
        fecha_seleccionada.insert(0, fecha_actual)  # Insertar la fecha actual
        self.label_load_analysis_data.config(text="")  # Limpiar texto en la etiqueta de archivo de analisis cargado
        self.label_load_study_data.config(text="")  # Limpiar texto en la etiqueta de archivo de estudio cargado
        self.FILE_PATH = ''
        self.STUDY_PATH = ''
        text_id_study.delete(0, tk.END)  # Limpiar entrada de ID del estudio

# -----------------------------------------
#             MÉTODOS VALIDACIÓN
# -----------------------------------------   

    # Método para validar que el usuario ha seleccionado algún valor en el combobox
    def validate_cmb(self, combo_1):
        # Lista de todos los Combobox
        comboboxes = [combo_1]
        
        for combo in comboboxes:
            if combo == '':
                # Si alguno de los Combobox no tiene un valor seleccionado, mostrar un mensaje de error
                messagebox.showerror("Error", "El desplegable debe tener un valor seleccionado")
                return False

        # Si todos los Combobox tienen un valor seleccionado, devolver True
        return True   
            

# -----------------------------------------
#        MÉTODOS MANIPULACIÓN DE DATOS
# ----------------------------------------- 

    # Método para extraer los datos específicos del estudio del que se va a crear el informe
    def filter_study_data(self, file_path, id_study):
        df = pd.read_csv(file_path, delimiter="\t")
        
        fila = df[df['Estudio.ID'] == int(id_study)]
        
        if not fila.empty:
            return fila
        else:
            messagebox.showinfo("Error", f"No se encontraron datos del estudio {id_study}.")                
            return None
        
    # Método para manipular datos del archivo de variantes
    def create_var_table(self, file_path):
        df = pd.read_csv(file_path, delimiter="\t")
        
        # Extraer Name, Condition(s), Variant type, Germline classification
        df_subset = df[['Name', 'Condition(s)', 'Germline classification']]
        
        # Reorganización de datos
        name_split = df_subset['Name'].str.split(" ", n=1, expand=True)
        df_subset['Cambio_proteico'] = name_split[1]
        
        name_split_2 = name_split[0].str.split(":", n=1, expand=True)
        df_subset['Cambio_nucleotidico'] = name_split_2[1]
        df_subset[['Transcrito', 'Gen']] = name_split_2[0].str.split("(", n=1, expand=True)
        df_subset['Gen'] = df_subset['Gen'].str.replace(")", "", regex=False)
        
        df_subset = df_subset.drop('Name', axis=1)
        
        return df_subset

    # Método para filtrar datos de variantes según lo seleccionado en el combobox
    def filter_var_table(self, df, combobox):
        # En el caso de que haya dos etiquetas, se selecciona la última
        df['Clasificacion'] = df['Germline classification'].apply(
            lambda efecto: efecto.split('/')[-1].strip() if isinstance(efecto, str) and '/' in efecto else efecto
        )
        
        dfs = {}
        if combobox == 'All':
            for clasificacion in df['Clasificacion'].unique():
                dfs[clasificacion] = df[df['Clasificacion'] == clasificacion]
        else:
            dfs[combobox] = df[df['Clasificacion'] == combobox]

        return dfs

    
# -----------------------------------------
#         MÉTODOS FORMATO INFORME
# -----------------------------------------         
    
    # Método para crear los diferentes estilos para el texto
    def create_text_style(self, doc):    
        # Crear estilo para el título general
        title_style = doc.styles.add_style('TitleStyle', WD_PARAGRAPH_ALIGNMENT.CENTER)
        title_font = title_style.font
        title_font.name = self.FONT_NAME
        title_font.size = Pt(13)
        title_font.bold = True
    
        # Crear estilo para el subtítulo general
        subtitle_style = doc.styles.add_style('SubtitleStyle', WD_PARAGRAPH_ALIGNMENT.CENTER)
        subtitle_font = subtitle_style.font
        subtitle_font.name = self.FONT_NAME
        subtitle_font.size = Pt(13)
    
        # Crear estilo para el título de los apartados
        parTitle_style = doc.styles.add_style('ParagraphTitleStyle', WD_PARAGRAPH_ALIGNMENT.CENTER)
        parTitle_font = parTitle_style.font
        parTitle_font.name = self.FONT_NAME
        parTitle_font.size = Pt(8)
        parTitle_font.bold = True
    
        # Crear estilo para texto simple
        parText_style = doc.styles.add_style('ParagraphTextStyle', WD_PARAGRAPH_ALIGNMENT.CENTER)
        parText_font = parText_style.font
        parText_font.name = self.FONT_NAME
        parText_font.size = Pt(9)       
    
    # Método para crear los diferentes estilos de celdas de tablas
    def create_cell_style(self, font, style):   
        font.name = self.FONT_NAME
        if style == 'table_TitleStyle':
            font.size = Pt(7.5)
            font.bold = True
        elif style == 'table_TextStyle':
            font.size = Pt(7.5)
            font.bold = False
        elif style == 'table_ConclusionTitleStyle':
            font.size = Pt(9)
            font.bold = True
        elif style == 'table_ConclusionTextStyle':
            font.size = Pt(9)
            font.bold = False
    
    # Método para añadir un párrafo al documento
    def add_report_paragraph(self, doc, texto, estilo=None, alineamiento=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, espacio_posterior=10):
        if estilo:
            p = doc.add_paragraph(style=estilo)
        else:
            p = doc.add_paragraph()

        # Ajustar el espaciado posterior
        p.paragraph_format.space_after = Pt(espacio_posterior)

        run = p.add_run(texto)

        if alineamiento:
            p.alignment = alineamiento    
    
    # Método para añadir un párrafo vacío al documento    
    def add_empty_paragraph(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    
    # Método para añadir un salto de página al documento    
    def add_report_page_break(self, doc):
        doc.add_page_break()
    
    # Método para añadir texto a una celda de tabla en el documento    
    def add_text_to_cell(self, cell, text, estilo=None, alineamiento=None):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = alineamiento
            for run in paragraph.runs:
                run.font.name = self.FONT_NAME
                # Aplicar el estilo si se proporciona
                if estilo:
                    self.create_cell_style(run.font, estilo)
    
    # Método para añadir una tabla al documento    
    def add_report_table(self, doc, df):
        df = df.drop('Germline classification', axis=1)
    
        nuevo_df = pd.DataFrame()
    
        for nrow, row in df.iterrows():
            nuevo_df = pd.concat([nuevo_df, pd.DataFrame([{
                'Clasificacion': row['Clasificacion'],
                'Gen': row['Gen'],
                'Transcrito': row['Transcrito'],
                'Cambio_nucleotidico': row['Cambio_nucleotidico'],
                'Cambio_proteico': row['Cambio_proteico']
            }])], ignore_index=True)
            nuevo_df = pd.concat([nuevo_df, pd.DataFrame([{'Gen': row['Condition(s)']}])], ignore_index=True)
            
        nuevo_df = nuevo_df.fillna("")
        
        # Agregar DataFrame como tabla
        table = doc.add_table(rows=nuevo_df.shape[0] + 1, cols=nuevo_df.shape[1])
        table.style = 'Colorful Shading Accent 5'
        for col_num, col_name in enumerate(nuevo_df.columns):
            self.add_text_to_cell(table.cell(0, col_num), col_name, estilo="table_TitleStyle", alineamiento=None)
            for row_num in range(nuevo_df.shape[0]):
                self.add_text_to_cell(table.cell(row_num + 1, col_num), str(nuevo_df.iloc[row_num, col_num]), estilo="table_TextStyle", alineamiento=None)
    
        # Combinar celdas para el texto de enfermedad cada dos líneas
        for nrow, row in enumerate(table.rows):
            if nrow > 1 and nrow % 2 == 0:
                table.cell(nrow, 1).merge(table.cell(nrow, 2)).merge(table.cell(nrow, 3)).merge(table.cell(nrow, 4))
                cell_text = table.cell(nrow, 1).text.split('\n')[0]
                self.add_text_to_cell(table.cell(nrow, 1), cell_text, estilo="table_TextStyle", alineamiento=None)
   
    # Método para añadir la cabecera a todas las páginas del documento    
    def add_report_header(self, doc, labels): 
        # Añadir una sección a todas las páginas
        for section in doc.sections:
            # Crear un nuevo encabezado para la sección
            header = section.header
            # Crear tabla
            table = header.add_table(rows=1, cols=2, width=Pt(500)) 
    
            # Ajustar tamaño de columnas
            table.autofit = False
    
            # Crear celdas de tabla
            cell1 = table.cell(0, 0)
            cell2 = table.cell(0, 1)

            # Agregar imagen a la izquierda
            cell1.paragraphs[0].add_run().add_picture(self.resource_path('images_folder\\logo.jpg'), width=Pt(50))
    
            # Ajustar tamaño de columnas
            for cell in table.columns[0].cells:
                cell.width = Inches(1.5)
                cell.height = Inches(0.2)

        return doc
    
    # Método que crea un elemento XML con un nombre dado
    def create_element(self, name):
        return OxmlElement(name)
    
    # Método que asigna un atributo con un nombre y valor especificados
    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)
    
    # Método que añade en el pie de página la versión y el número de página sobre el total
    def add_page_number(self, paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.style = 'ParagraphTextStyle'
    
        # Agregar versión
        paragraph.add_run(f"{self.APP_VERSION}\t\t\t\t\t\t\t\t\t\t\t")  # Puedes ajustar el formato según tus necesidades
    
        page_num_run = paragraph.add_run()
    
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
    
        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')
    
        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)
    
        of_run = paragraph.add_run()
        t2 = self.create_element('w:t')
        self.create_attribute(t2, 'xml:space', 'preserve')
        t2.text = '/'
        of_run._r.append(t2)
    
        fldChar3 = self.create_element('w:fldChar')
        self.create_attribute(fldChar3, 'w:fldCharType', 'begin')
    
        instrText2 = self.create_element('w:instrText')
        self.create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"
    
        fldChar4 = self.create_element('w:fldChar')
        self.create_attribute(fldChar4, 'w:fldCharType', 'end')
    
        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)    


# -----------------------------------------
#       MÉTODO CREACIÓN DE INFORME
# -----------------------------------------     
    
    # Método que genera el informe y lo exporta
    def generate_report(self, combo_1, fecha_seleccionada, id_study):

        # Establecer la configuración regional en español
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
        
        # Validar que se han introducido todos los datos de entrada
        if not self.validate_cmb(combo_1):
            return

        # Validar que se ha cargado el archivo de análisis
        if not self.label_load_analysis_data.cget("text"):
            messagebox.showerror("Error", "No se ha cargado ningún archivo.")
            return
        
        # Manipular datos de variantes
        df_total = self.create_var_table(self.FILE_PATH)
        df_filtrado = self.filter_var_table(df_total, combo_1)
        
        # Extraer datos de estudio
        datos_estudio = self.filter_study_data(self.STUDY_PATH, id_study)
        if datos_estudio is None or datos_estudio.empty:
            return
        
        # Cargar las etiquetas desde el archivo labels.py
        labels = self.load_labels(self.PATH_LABELS)  

        # Crear documento word - informe
        doc = Document()

        # Crear estilos
        self.create_text_style(doc)
    
        # Inicializar variables globales
        n_tabla = 0
        
        ##### ENCABEZADO #####
    
        # Agregar encabezado a todas las páginas
        self.add_report_header(doc, labels)

    
        ##### TÍTULO Y SUBTÍTULO #####
    
        # Agregar título
        self.add_report_paragraph(doc, labels['titulo'], estilo='TitleStyle', alineamiento = WD_PARAGRAPH_ALIGNMENT.CENTER)

        # Agregar subtítulo
        self.add_report_paragraph(doc, labels['subtitulo'].format(datos_estudio['Enfermedad/Gen'].values[0]), estilo='SubtitleStyle', alineamiento = WD_PARAGRAPH_ALIGNMENT.CENTER)  #.format(combo_2).upper()

        
        ##### DATOS DEL ESTUDIO #####
        
        # Agregar cabecera 
        self.add_report_paragraph(doc, labels['titulo_datos_estudio'], estilo = 'ParagraphTitleStyle')
    
        # Crear tabla
        tabla_datos_demo = doc.add_table(rows=2, cols=2, style = 'Light List Accent 5')
        tabla_datos_demo.autofit = False
        for cell in tabla_datos_demo.columns[0].cells:
            cell.width = Inches(2.4)
        for cell in tabla_datos_demo.columns[1].cells:
            cell.width = Inches(3.6)

        # Crear celdas de tabla
        cell1 = tabla_datos_demo.cell(0, 0)
        cell2 = tabla_datos_demo.cell(0, 1)
        cell3 = tabla_datos_demo.cell(1, 0)
        cell4 = tabla_datos_demo.cell(1, 1)

        # Agregar texto a cada celda
        self.add_text_to_cell(cell1, labels['informe_datos_estudio_id'].format(datos_estudio['Estudio.ID'].values[0]), estilo="table_ConclusionTextStyle")
        self.add_text_to_cell(cell2, labels['informe_datos_estudio_fecha'].format(fecha_seleccionada), estilo = "table_ConclusionTextStyle")
        self.add_text_to_cell(cell3, labels['informe_datos_estudio_objetivo'].format(datos_estudio['Objetivo'].values[0]), estilo="table_ConclusionTextStyle")
        self.add_text_to_cell(cell4, labels['informe_datos_estudio_filtrado'].format(datos_estudio['Filtro'].values[0]), estilo = "table_ConclusionTextStyle")
    
        ##### LÍNEA VACÍA #####
    
        self.add_empty_paragraph(doc)    


        ##### RESULTADOS DEL ANALISIS #####

        # Añadir tablas con datos + párrafo asociado
        numero_variantes = 0
        datos_variantes = []
        
        all_empty = all(df.empty for df in df_filtrado.values())
        
        if not all_empty:
            # Agregar cabecera 
            self.add_report_paragraph(doc, labels['titulo_resultados'], estilo = 'ParagraphTitleStyle')
            
            for key, df in df_filtrado.items():
                numero_variantes += len(df)

                self.add_report_table(doc, df)
                doc.add_paragraph()
                clasificacion = key.replace(" ", "_")
                datos_variantes.append(key)

                for nrow, row in df.iterrows():
                    self.add_report_paragraph(doc, labels[f"texto_{clasificacion}"].format(df['Gen'][nrow], df['Cambio_nucleotidico'][nrow], df['Transcrito'][nrow], df['Cambio_proteico'][nrow]), estilo='ParagraphTextStyle')
            
            
        ##### CONCLUSION #####
        
        no_datos_variantes = ['Conflicting classifications', 'Benign', 'Likely benign', 'Uncertain significance', 'Likely pathogenic', 'Pathogenic']
        differences = set(no_datos_variantes) - set(datos_variantes)

        # Crear tabla
        tabla_conclusiones = doc.add_table(rows=3, cols=1, style = 'Medium List 2') 

        # Crear celdas
        cell_c = tabla_conclusiones.cell(0, 0)
        cell_t = tabla_conclusiones.cell(1, 0)
        cell_t_2 = tabla_conclusiones.cell(2, 0)

        # Agregar cabecera
        self.add_text_to_cell(cell_c, labels['titulo_conclusion'], estilo = "table_ConclusionTitleStyle")
        
        # Añadir texto conclusion
        if numero_variantes == 0:
            self.add_text_to_cell(cell_t, labels['tabla_vacia'].format(combo_1), estilo = "table_ConclusionTextStyle")
        else:
            datos_variantes = str(datos_variantes).replace("[", "").replace("]", "")
            self.add_text_to_cell(cell_t,  labels['conclusion_general'].format(numero_variantes, datos_variantes), estilo = "table_ConclusionTextStyle")
            
            if combo_1 == 'All' and no_datos_variantes:
                no_datos_variantes = str(no_datos_variantes).replace("[", "").replace("]", "")
                self.add_text_to_cell(cell_t_2,  labels['conclusion_alguna_tabla_vacia'].format(differences).replace("{", "").replace("}", ""), estilo = "table_ConclusionTextStyle")

        
        ##### PIE DE PÁGINA #####
        
        self.add_page_number(doc.sections[0].footer.paragraphs[0])            
                
        # Seleccionar ruta para guardar archivo
        output_path = asksaveasfile(mode = 'w', defaultextension = ".docx", filetypes = [("Archivos de Word", "*.docx")])
        # Verificar si se introdujo una ruta antes de guardar
        if output_path:
            doc.save(output_path.name)
            os.startfile(output_path.name)
            output_path.close()
            messagebox.showinfo("Info", "Se ha generado el informe correctamente.")
        else:
            messagebox.showinfo("Info", "No se ha indicado una ruta para guardar.")
    

if __name__ == "__main__":
    locale.setlocale(locale.LC_TIME, 'es_ES')
    root = tk.Tk()
    app = InterfazApp(root)
